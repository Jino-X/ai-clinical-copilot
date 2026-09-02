from __future__ import annotations

import time
import uuid
from typing import Annotated
from uuid import UUID

from fastapi import APIRouter, Query, Request, status

from app.api.deps import AuditDep, DatabaseDep, OrganizationDep, TenantConnection
from app.core.errors import NotFoundError, ServiceUnavailableError
from app.core.logging import get_logger
from app.core.permissions import Permission
from app.providers.factory import ProviderFactory
from app.repositories.ai_generations import AiGenerationRepository
from app.repositories.documents import DocumentRepository
from app.repositories.patients import PatientRepository
from app.schemas.documents import (
    CreateUploadUrlRequest,
    CreateUploadUrlResponse,
    DocumentCategory,
    DocumentDownloadUrlResponse,
    DocumentStatus,
    MedicalDocumentResponse,
    MedicalDocumentSummary,
    UpdateDocumentRequest,
    VerifyDocumentRequest,
)
from app.services.ai.extraction import DocumentExtractionService
from app.services.audit.service import AuditAction
from app.services.storage.service import StorageService

router = APIRouter(prefix="/documents", tags=["documents"])

_repo = DocumentRepository()
_patient_repo = PatientRepository()

logger = get_logger(__name__)
_ai_repo = AiGenerationRepository()


def _get_storage_service(request: Request) -> StorageService:
    service: StorageService = request.app.state.storage_service
    return service


def _get_provider_factory(request: Request) -> ProviderFactory:
    factory: ProviderFactory = request.app.state.provider_factory
    return factory


# ===========================================================================
# Upload
# ===========================================================================


@router.post(
    "/upload-url",
    response_model=CreateUploadUrlResponse,
    summary="Get a signed URL for uploading a document",
)
async def create_upload_url(
    payload: CreateUploadUrlRequest,
    context: OrganizationDep,
    connection: TenantConnection,
    audit: AuditDep,
    request: Request,
) -> CreateUploadUrlResponse:
    """Get a signed URL for uploading a medical document.

    The client uploads directly to Supabase Storage via the signed URL. The
    document is stored in a private bucket and is never publicly accessible
    (PRD §9).
    """
    context.require(Permission.PATIENT_WRITE)

    patient = await _patient_repo.get(connection, patient_id=payload.patient_id)
    if patient is None or patient.organization_id != context.organization_id:
        raise NotFoundError("Patient not found")

    storage = _get_storage_service(request)

    # Create the document record first so we have an ID for the storage path.
    # We generate the UUID here so the storage_path (which includes the ID)
    # matches the DB row ID exactly.
    document_id = uuid.uuid4()
    storage_path = storage._document_storage_path(
        str(context.organization_id),
        str(document_id),
        payload.file_name,
    )
    full_storage_path = f"medical-documents/{storage_path}"

    # Create the document record with the pre-generated ID.
    doc = await _repo.create(
        connection,
        organization_id=context.organization_id,
        patient_id=payload.patient_id,
        uploaded_by=context.user.id,
        title=payload.title,
        storage_path=full_storage_path,
        file_name=payload.file_name,
        content_type=payload.content_type,
        file_size_bytes=payload.file_size_bytes,
        document_id=document_id,
    )

    # Now create the signed upload URL.
    signed = await storage.create_document_upload_url(
        organization_id=str(context.organization_id),
        document_id=str(doc.id),
        file_name=payload.file_name,
        content_type=payload.content_type,
    )

    await audit.record(
        AuditAction.DOCUMENT_UPLOADED,
        actor_user_id=context.user.id,
        organization_id=context.organization_id,
        resource_type="document",
        resource_id=str(doc.id),
        request=request,
        metadata={"patient_id": str(payload.patient_id)},
    )

    return CreateUploadUrlResponse(
        upload_url=signed.upload_url,
        storage_path=signed.storage_path,
        document_id=doc.id,
        expires_at=signed.expires_at,
    )


# ===========================================================================
# List and get
# ===========================================================================


@router.get(
    "",
    response_model=list[MedicalDocumentSummary],
    summary="List documents",
)
async def list_documents(
    context: OrganizationDep,
    connection: TenantConnection,
    patient_id: Annotated[UUID | None, Query()] = None,
    limit: Annotated[int, Query(ge=1, le=100)] = 50,
    offset: Annotated[int, Query(ge=0)] = 0,
) -> list[MedicalDocumentSummary]:
    context.require(Permission.PATIENT_READ)
    if patient_id is not None:
        return await _repo.list_for_patient(
            connection,
            patient_id=patient_id,
            limit=limit,
            offset=offset,
        )
    return await _repo.list_for_organization(
        connection,
        organization_id=context.organization_id,
        limit=limit,
        offset=offset,
    )


@router.get(
    "/{document_id}",
    response_model=MedicalDocumentResponse,
    summary="Get a document",
)
async def get_document(
    document_id: UUID,
    context: OrganizationDep,
    connection: TenantConnection,
    audit: AuditDep,
    request: Request,
) -> MedicalDocumentResponse:
    context.require(Permission.PATIENT_READ)
    doc = await _repo.get(connection, document_id=document_id)
    if doc is None:
        raise NotFoundError("Document not found")
    await audit.record(
        AuditAction.DOCUMENT_VIEWED,
        actor_user_id=context.user.id,
        organization_id=context.organization_id,
        resource_type="document",
        resource_id=str(document_id),
        request=request,
    )
    return doc


@router.delete(
    "/{document_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete a document",
)
async def delete_document(
    document_id: UUID,
    context: OrganizationDep,
    connection: TenantConnection,
    audit: AuditDep,
    request: Request,
) -> None:
    context.require(Permission.PATIENT_WRITE)
    storage_path = await _repo.delete(connection, document_id=document_id)
    if storage_path is None:
        raise NotFoundError("Document not found")

    # Best-effort storage cleanup — don't fail the request if the object
    # is already gone or storage is not configured.
    storage = _get_storage_service(request)
    try:
        await storage.delete_object(storage_path=storage_path)
    except Exception:
        # Log only the exception type, never the path (may contain PHI).
        logger.warning("document_storage_delete_failed", error_type="storage")

    await audit.record(
        AuditAction.DOCUMENT_DELETED,
        actor_user_id=context.user.id,
        organization_id=context.organization_id,
        resource_type="document",
        resource_id=str(document_id),
        request=request,
    )


# ===========================================================================
# Download URL
# ===========================================================================


@router.get(
    "/{document_id}/download-url",
    response_model=DocumentDownloadUrlResponse,
    summary="Get a signed URL for downloading a document",
)
async def get_download_url(
    document_id: UUID,
    context: OrganizationDep,
    connection: TenantConnection,
    request: Request,
) -> DocumentDownloadUrlResponse:
    context.require(Permission.PATIENT_READ)
    doc = await _repo.get(connection, document_id=document_id)
    if doc is None:
        raise NotFoundError("Document not found")

    storage = _get_storage_service(request)
    signed = await storage.create_download_url(storage_path=doc.storage_path)
    return DocumentDownloadUrlResponse(
        download_url=signed.download_url,
        expires_at=signed.expires_at,
        content_type=signed.content_type,
        size_bytes=signed.size_bytes,
    )


# ===========================================================================
# Extract (OCR + classification + medical extraction)
# ===========================================================================


@router.post(
    "/{document_id}/extract",
    response_model=MedicalDocumentResponse,
    summary="Extract and classify document content",
)
async def extract_document(
    document_id: UUID,
    context: OrganizationDep,
    connection: TenantConnection,
    database: DatabaseDep,
    audit: AuditDep,
    request: Request,
) -> MedicalDocumentResponse:
    """Extract text from the document, classify it, and extract medical info.

    Downloads the document from Storage, extracts text (for text-based
    documents), and uses the LLM to classify and extract structured medical
    information. The extracted data is a draft requiring doctor verification
    (PRD §9, §12). If the AI fails, the document is safe (PRD §24).
    """
    context.require(Permission.PATIENT_WRITE)

    doc = await _repo.get(connection, document_id=document_id)
    if doc is None:
        raise NotFoundError("Document not found")

    factory = _get_provider_factory(request)
    if not factory.llm_configured:
        raise ServiceUnavailableError("LLM provider is not configured")

    storage = _get_storage_service(request)

    # Mark as processing.
    await _repo.update_status(
        connection,
        document_id=document_id,
        status=DocumentStatus.PROCESSING,
    )

    gen_id = await _ai_repo.create(
        connection,
        organization_id=context.organization_id,
        task_type="document_extraction",
        provider=factory.llm.name,
        model=factory.llm.model,
    )

    start_time = time.monotonic()

    try:
        # Download the document.
        content_bytes, content_type = await storage.download_document(
            storage_path=doc.storage_path
        )

        # Extract text. For text-based formats (PDF, DOCX, TXT), we extract
        # directly. For images (JPG, PNG), OCR would be needed. For the MVP,
        # we decode text-based content; images require an OCR provider.
        extracted_text = _extract_text(content_bytes, content_type or doc.content_type)
        # PostgreSQL text columns reject null bytes (0x00); strip them.
        extracted_text = extracted_text.replace("\x00", "")

        if not extracted_text or not extracted_text.strip():
            # Could not extract text (e.g., image without OCR).
            await _repo.update_status(
                connection,
                document_id=document_id,
                status=DocumentStatus.FAILED,
                error_message="Could not extract text from document",
            )
            raise ServiceUnavailableError(
                "Could not extract text from this document type"
            )

        # Classify and extract medical information via LLM.
        service = DocumentExtractionService(factory.llm)
        result = await service.extract(document_text=extracted_text)

        duration_ms = int((time.monotonic() - start_time) * 1000)

        # Parse the category.
        category_str = result.get("category", "other")
        try:
            category = DocumentCategory(category_str)
        except ValueError:
            category = DocumentCategory.OTHER

        updated = await _repo.update_extraction(
            connection,
            document_id=document_id,
            extracted_text=extracted_text,
            extracted_data=result,
            category=category,
            provider=factory.llm.name,
            model=factory.llm.model,
        )

        async with database.privileged() as priv_conn:
            await _ai_repo.mark_completed(
                priv_conn, generation_id=gen_id, duration_ms=duration_ms
            )

        await audit.record(
            AuditAction.AI_NOTE_GENERATED,
            actor_user_id=context.user.id,
            organization_id=context.organization_id,
            resource_type="document",
            resource_id=str(document_id),
            request=request,
            metadata={"task": "document_extraction"},
        )

        if updated is None:
            raise NotFoundError("Document not found")
        return updated

    except ServiceUnavailableError:
        raise
    except Exception as exc:
        await _repo.update_status(
            connection,
            document_id=document_id,
            status=DocumentStatus.FAILED,
            error_message=str(exc)[:500],
        )
        async with database.privileged() as priv_conn:
            await _ai_repo.mark_failed(
                priv_conn,
                generation_id=gen_id,
                error_message=str(exc)[:500],
            )
        raise


# ===========================================================================
# Update and verify
# ===========================================================================


@router.patch(
    "/{document_id}",
    response_model=MedicalDocumentResponse,
    summary="Update document metadata",
)
async def update_document(
    document_id: UUID,
    payload: UpdateDocumentRequest,
    context: OrganizationDep,
    connection: TenantConnection,
) -> MedicalDocumentResponse:
    context.require(Permission.PATIENT_WRITE)
    doc = await _repo.update(
        connection,
        document_id=document_id,
        title=payload.title,
        category=payload.category,
    )
    if doc is None:
        raise NotFoundError("Document not found")
    return doc


@router.post(
    "/{document_id}/verify",
    response_model=MedicalDocumentResponse,
    summary="Verify extracted document information (PRD §9)",
)
async def verify_document(
    document_id: UUID,
    payload: VerifyDocumentRequest,
    context: OrganizationDep,
    connection: TenantConnection,
    audit: AuditDep,
    request: Request,
) -> MedicalDocumentResponse:
    """Doctor verifies the extracted information.

    The doctor can edit the extracted data and category before verifying.
    Once verified, the document's status becomes 'verified' and a timeline
    event is created (PRD §9).
    """
    context.require(Permission.PATIENT_WRITE)

    doc = await _repo.get(connection, document_id=document_id)
    if doc is None:
        raise NotFoundError("Document not found")
    if doc.status == DocumentStatus.VERIFIED:
        from app.core.errors import ConflictError

        raise ConflictError("Document is already verified")

    verified = await _repo.verify(
        connection,
        document_id=document_id,
        verified_by=context.user.id,
        category=payload.category,
        extracted_data=payload.extracted_data,
    )
    if verified is None:
        raise NotFoundError("Document not found")

    await audit.record(
        AuditAction.DOCUMENT_VIEWED,
        actor_user_id=context.user.id,
        organization_id=context.organization_id,
        resource_type="document",
        resource_id=str(document_id),
        request=request,
        metadata={"action": "verified"},
    )

    return verified


# ===========================================================================
# Helpers
# ===========================================================================


def _extract_text(content: bytes, content_type: str) -> str:
    """Extract text from document content based on content type.

    Handles TXT, JSON, HTML, XML, DOCX (via python-docx), and PDF (via pypdf).
    Image-based documents (JPG, PNG) require OCR and are not supported yet.
    """
    ct = content_type.lower()

    # Plain text.
    if "text/plain" in ct:
        return content.decode("utf-8", errors="replace").replace("\x00", "")

    # JSON (some APIs return structured data as JSON).
    if "application/json" in ct:
        return content.decode("utf-8", errors="replace").replace("\x00", "")

    # HTML / XML — but NOT Office Open XML (DOCX), which is a ZIP.
    if ("text" in ct or "html" in ct) and "officedocument" not in ct:
        return content.decode("utf-8", errors="replace").replace("\x00", "")

    # DOCX — use python-docx for reliable text extraction.
    if "wordprocessingml" in ct or "officedocument.wordprocessing" in ct:
        return _extract_docx_text(content)

    # PDF — use pypdf for text extraction.
    if "pdf" in ct:
        return _extract_pdf_text(content)

    # Image formats require OCR — not supported yet.
    if "image" in ct:
        return ""

    # Unknown binary format — try DOCX first (most common), then PDF.
    text = _extract_docx_text(content)
    if text:
        return text
    return _extract_pdf_text(content)


def _extract_docx_text(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    import io

    from docx import Document

    try:
        doc = Document(io.BytesIO(content))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract text from tables.
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs).replace("\x00", "")
    except Exception:
        return ""


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    import io

    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages).replace("\x00", "")
    except Exception:
        return ""
